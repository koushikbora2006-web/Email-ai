import io
import os

class ExportService:
    @staticmethod
    def generate_docx(subject, body, sender="Email AI Assistant"):
        """Generate DOCX document for email download."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            
            doc = Document()
            
            # Header
            header_p = doc.add_paragraph()
            r = header_p.add_run("Email Writer AI Draft")
            r.font.name = 'Arial'
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(128, 128, 128)
            
            # Title / Subject
            h1 = doc.add_heading(level=1)
            h1_run = h1.add_run(f"Subject: {subject}")
            h1_run.font.name = 'Arial'
            h1_run.font.size = Pt(16)
            h1_run.font.bold = True
            h1_run.font.color.rgb = RGBColor(30, 41, 59)
            
            doc.add_paragraph() # spacing
            
            # Body
            for paragraph in body.split('\n'):
                if paragraph.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(paragraph.strip())
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(51, 65, 85)
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            print(f"[DOCX Export Error]: {e}")
            # Fallback basic text buffer
            return f"Subject: {subject}\n\n{body}".encode('utf-8')

    @staticmethod
    def generate_pdf(subject, body, sender="Email AI Assistant"):
        """Generate PDF document for email download using ReportLab."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib import colors
            
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=54,
                leftMargin=54,
                topMargin=54,
                bottomMargin=54
            )
            
            styles = getSampleStyleSheet()
            
            title_style = ParagraphStyle(
                'EmailSubject',
                parent=styles['Heading1'],
                fontName='Helvetica-Bold',
                fontSize=16,
                textColor=colors.HexColor('#1E293B'),
                spaceAfter=15
            )
            
            body_style = ParagraphStyle(
                'EmailBody',
                parent=styles['Normal'],
                fontName='Helvetica',
                fontSize=11,
                leading=16,
                textColor=colors.HexColor('#334155'),
                spaceAfter=10
            )
            
            story = []
            story.append(Paragraph(f"<b>Subject:</b> {subject}", title_style))
            story.append(Spacer(1, 10))
            
            for line in body.split('\n'):
                if line.strip():
                    # Escape HTML characters in PDF text
                    clean_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(clean_line, body_style))
                else:
                    story.append(Spacer(1, 8))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            print(f"[PDF Export Error]: {e}")
            return f"Subject: {subject}\n\n{body}".encode('utf-8')
